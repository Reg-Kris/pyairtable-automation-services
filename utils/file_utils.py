"""
File content extraction utilities for PyAirtable Automation Services.
"""

import os
import hashlib
import mimetypes
from typing import Optional, Dict, Any, Union
import logging
from pathlib import Path
import asyncio
import subprocess

logger = logging.getLogger(__name__)


class FileExtractor:
    """File content extraction utilities."""
    
    @staticmethod
    def get_file_hash(file_path: str) -> str:
        """Calculate SHA-256 hash of file."""
        sha256_hash = hashlib.sha256()
        try:
            with open(file_path, "rb") as f:
                for byte_block in iter(lambda: f.read(4096), b""):
                    sha256_hash.update(byte_block)
            return sha256_hash.hexdigest()
        except Exception as e:
            logger.error(f"Error calculating file hash for {file_path}: {e}")
            return ""
    
    @staticmethod
    def get_mime_type(file_path: str) -> str:
        """Get MIME type of file."""
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type or "application/octet-stream"
    
    @staticmethod
    def get_file_extension(filename: str) -> str:
        """Get file extension in lowercase."""
        return Path(filename).suffix.lower()
    
    @staticmethod
    async def extract_text_content(file_path: str) -> Dict[str, Any]:
        """Extract text content from various file types."""
        if not os.path.exists(file_path):
            return {"error": "File not found", "content": "", "metadata": {}}
        
        extension = FileExtractor.get_file_extension(file_path)
        
        try:
            if extension == ".txt":
                return await FileExtractor._extract_txt(file_path)
            elif extension == ".pdf":
                return await FileExtractor._extract_pdf(file_path)
            elif extension in [".doc", ".docx"]:
                return await FileExtractor._extract_docx(file_path)
            elif extension in [".csv"]:
                return await FileExtractor._extract_csv(file_path)
            elif extension in [".xlsx", ".xls"]:
                return await FileExtractor._extract_excel(file_path)
            else:
                return {
                    "error": f"Unsupported file type: {extension}",
                    "content": "",
                    "metadata": {"file_type": extension}
                }
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return {
                "error": f"Extraction failed: {str(e)}",
                "content": "",
                "metadata": {"file_type": extension}
            }
    
    @staticmethod
    async def _extract_txt(file_path: str) -> Dict[str, Any]:
        """Extract content from text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            content = ""
            encoding_used = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    encoding_used = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if not encoding_used:
                return {
                    "error": "Could not decode text file",
                    "content": "",
                    "metadata": {"file_type": ".txt"}
                }
            
            return {
                "content": content,
                "metadata": {
                    "file_type": ".txt",
                    "encoding": encoding_used,
                    "character_count": len(content),
                    "line_count": content.count('\n') + 1 if content else 0
                }
            }
        except Exception as e:
            return {
                "error": f"Text extraction failed: {str(e)}",
                "content": "",
                "metadata": {"file_type": ".txt"}
            }
    
    @staticmethod
    async def _extract_pdf(file_path: str) -> Dict[str, Any]:
        """Extract content from PDF file."""
        try:
            import PyPDF2
            
            content = ""
            metadata = {"file_type": ".pdf", "pages": 0}
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata["pages"] = len(pdf_reader.pages)
                
                # Limit pages to prevent memory issues
                max_pages = min(len(pdf_reader.pages), 1000)
                
                for page_num in range(max_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
                        continue
                
                # Extract PDF metadata
                if pdf_reader.metadata:
                    metadata.update({
                        "title": pdf_reader.metadata.get("/Title", ""),
                        "author": pdf_reader.metadata.get("/Author", ""),
                        "subject": pdf_reader.metadata.get("/Subject", ""),
                        "creator": pdf_reader.metadata.get("/Creator", ""),
                    })
            
            return {
                "content": content.strip(),
                "metadata": metadata
            }
        except ImportError:
            return {
                "error": "PyPDF2 not installed",
                "content": "",
                "metadata": {"file_type": ".pdf"}
            }
        except Exception as e:
            return {
                "error": f"PDF extraction failed: {str(e)}",
                "content": "",
                "metadata": {"file_type": ".pdf"}
            }
    
    @staticmethod
    async def _extract_docx(file_path: str) -> Dict[str, Any]:
        """Extract content from DOCX file."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    content += " | ".join(row_text) + "\n"
            
            metadata = {
                "file_type": ".docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "character_count": len(content)
            }
            
            # Extract document properties
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            
            return {
                "content": content.strip(),
                "metadata": metadata
            }
        except ImportError:
            return {
                "error": "python-docx not installed",
                "content": "",
                "metadata": {"file_type": ".docx"}
            }
        except Exception as e:
            return {
                "error": f"DOCX extraction failed: {str(e)}",
                "content": "",
                "metadata": {"file_type": ".docx"}
            }
    
    @staticmethod
    async def _extract_csv(file_path: str) -> Dict[str, Any]:
        """Extract content from CSV file."""
        try:
            import pandas as pd
            
            # Try different encodings and separators
            encodings = ['utf-8', 'latin-1', 'cp1252']
            separators = [',', ';', '\t']
            
            df = None
            encoding_used = None
            separator_used = None
            
            for encoding in encodings:
                for sep in separators:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding, sep=sep, nrows=1000)  # Limit rows
                        if len(df.columns) > 1:  # Good separator found
                            encoding_used = encoding
                            separator_used = sep
                            break
                    except:
                        continue
                if df is not None and len(df.columns) > 1:
                    break
            
            if df is None:
                # Fallback to default
                df = pd.read_csv(file_path, nrows=1000)
                encoding_used = 'utf-8'
                separator_used = ','
            
            # Convert to string representation
            content = df.to_string(index=False)
            
            metadata = {
                "file_type": ".csv",
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": list(df.columns),
                "encoding": encoding_used,
                "separator": separator_used
            }
            
            return {
                "content": content,
                "metadata": metadata
            }
        except ImportError:
            return {
                "error": "pandas not installed",
                "content": "",
                "metadata": {"file_type": ".csv"}
            }
        except Exception as e:
            return {
                "error": f"CSV extraction failed: {str(e)}",
                "content": "",
                "metadata": {"file_type": ".csv"}
            }
    
    @staticmethod
    async def _extract_excel(file_path: str) -> Dict[str, Any]:
        """Extract content from Excel file."""
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            content = ""
            sheet_info = []
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name, nrows=1000)
                    content += f"\n--- Sheet: {sheet_name} ---\n"
                    content += df.to_string(index=False) + "\n"
                    
                    sheet_info.append({
                        "name": sheet_name,
                        "rows": len(df),
                        "columns": len(df.columns),
                        "column_names": list(df.columns)
                    })
                except Exception as e:
                    logger.warning(f"Error reading sheet {sheet_name}: {e}")
                    continue
            
            metadata = {
                "file_type": Path(file_path).suffix.lower(),
                "sheet_count": len(excel_file.sheet_names),
                "sheets": sheet_info
            }
            
            return {
                "content": content.strip(),
                "metadata": metadata
            }
        except ImportError:
            return {
                "error": "pandas and openpyxl not installed",
                "content": "",
                "metadata": {"file_type": ".xlsx"}
            }
        except Exception as e:
            return {
                "error": f"Excel extraction failed: {str(e)}",
                "content": "",
                "metadata": {"file_type": ".xlsx"}
            }
    
    @staticmethod
    def validate_file_size(file_path: str, max_size: int) -> bool:
        """Validate file size."""
        try:
            return os.path.getsize(file_path) <= max_size
        except Exception:
            return False
    
    @staticmethod
    def cleanup_file(file_path: str) -> bool:
        """Delete file safely."""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
            return False
        except Exception as e:
            logger.error(f"Error deleting file {file_path}: {e}")
            return False